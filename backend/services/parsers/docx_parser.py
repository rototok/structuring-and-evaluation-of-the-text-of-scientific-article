from collections import Counter
from typing import List

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.text.run import Run
from docx.oxml.ns import qn


class DocumentParser:
    def __init__(self, file_name: str):
        self.document = Document(file_name)


    def get_tables(self) -> List[Table]:
        return self.document.tables

    def get_paragraphs(self) -> List[Paragraph]:
        return self.document.paragraphs
    
    def get_blocks(self) -> List:
        return self.document.element.body
    
    def get_runs(self, paragraph: Paragraph) -> List[Run]:
        return paragraph.runs
    
    def get_text(self, paragraph: Paragraph) -> str:
        return paragraph.text


    def is_paragraph(self, block) -> bool:
        if block.tag.endswith('p'):
            return True
        
        return False
    
    def is_table(self, block) -> bool:
        if block.tag.endswith('tbl'):
            return True
        
        return False
    
    def is_bold(self, paragraph: Paragraph) -> bool:
        runs = self.get_runs(paragraph)
        counter = 0

        if not runs:
            return False
        
        for run in runs:
            if run.bold:
                counter += 1
        
        if counter == len(runs):
            return True
        
        return False


    def get_outline_level(self, paragraph) -> int | None:
        outline_lvl = self.__get_outline_from_element(paragraph._element)

        if outline_lvl is None:
            outline_lvl = self.__get_outline_from_style(paragraph.style)
        
        return outline_lvl
    
    def get_font_sizes(self, paragraph) -> List[float]:
        font_sizes = []
        for run in paragraph.runs:
            if run.font.size:
                font_sizes.append(run.font.size.pt)

        return font_sizes    

    def __get_outline_from_style(self, style) -> int | None:
        visited = set()

        while style is not None and style.style_id not in visited:
            visited.add(style.style_id)

            outline_lvl = self.__get_outline_from_element(style._element)
            
            if outline_lvl is not None:
                return outline_lvl

            style = style.base_style

        return None

    def __get_outline_from_element(self, element) -> int | None:
        pPr = element.find(qn("w:pPr"))
        if pPr is None:
            return None
        
        outline = pPr.find(qn("w:outlineLvl"))
            
        if outline is None:
            return None
        
        outline_lvl = int(outline.get(qn("w:val")))
        return outline_lvl


class DocumentAnalyzer:
    def __init__(self):
        self.__font_size_counter = Counter()
        self.__font_levels = dict()

    @property
    def font_size_counter(self):
        return self.__font_size_counter
    
    @property
    def font_levels(self):
        return self.__font_levels
    
    @font_levels.setter
    def font_levels(self, font_levels):
        self.__font_levels = font_levels
    
    def count_font_sizes(self, document_parser: DocumentParser) -> Counter:
        paragraphs = document_parser.get_paragraphs()

        for paragraph in paragraphs:
            font_sizes = document_parser.get_font_sizes(paragraph)
            self.font_size_counter.update(font_sizes)
    
        return self.font_size_counter

    def determine_font_levels(self) -> dict | None:
        common_size = self.font_size_counter.most_common(1)[0][0]

        heading_font_sizes = [font_size for font_size in self.font_size_counter.items() if font_size[0] > common_size]
        
        if len(heading_font_sizes) == 0:
            return None
        
        heading_font_sizes = sorted(heading_font_sizes, reverse=True)
        
        font_levels = dict()
        for font_level, font_size in enumerate(heading_font_sizes, start=1):
            font_levels[font_size] = font_level

        self.font_levels = font_levels

        return font_levels

    def count_heading_score(self, paragraph: Paragraph, document_parser: DocumentParser) -> int:
        score = 0

        paragraph_text = document_parser.get_text(paragraph)
        paragraph_text_len = len(paragraph_text.split())

        if document_parser.is_bold(paragraph):
            score += 1

        if paragraph_text_len < 20 and paragraph_text_len > 0:
            score += 1

        if not paragraph_text.endswith("."):
            score += 2
        
        if paragraph_text.isupper():
            score += 2

        return score

class MarkdownConverter:
    def __init__(self, document_parser: DocumentParser):
        self.document_parser = document_parser

    def convert_to_markdown(self, document_statistics: DocumentAnalyzer) -> str:
        converted_text = []

        paragraphs = self.document_parser.get_paragraphs()
        tables = self.document_parser.get_tables()

        document_statistics.count_font_sizes(self.document_parser)
        document_statistics.determine_font_levels()

        for block in self.document_parser.get_blocks():
            if self.document_parser.is_paragraph(block):
                paragraph_text = ""

                paragraph = paragraphs.pop(0)
                heading_lvl = self.detect_heading_level(paragraph, document_statistics)

                if heading_lvl is not None:
                    paragraph_text += self.heading_prefix(heading_lvl)
                
                paragraph_text += self.convert_paragraph(paragraph)
                converted_text.append(paragraph_text)

            elif self.document_parser.is_table(block):
                table = tables.pop(0)
                table_text = self.convert_table(table)

                converted_text.append(table_text)
        
        return "\n\n".join(converted_text)
    
    def convert_table(self, table: Table) -> str:
        table_text = ""

        for i, row in enumerate(table.rows):
            table_text += "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |\n"
            if i == 0:
                table_text += "| " + " | ".join("---" for _ in row.cells) + " |\n"

        return table_text

    def detect_heading_level(self, paragraph: Paragraph, document_statistics: DocumentAnalyzer) -> int | None:
        outline_lvl = self.document_parser.get_outline_level(paragraph)
        if outline_lvl is not None:
            return outline_lvl + 1
        
        font_sizes = self.document_parser.get_font_sizes(paragraph)
        if len(font_sizes) == 1:
            if font_lvl := document_statistics.font_levels.get(font_sizes[0]):
                return font_lvl
            
        heading_score = document_statistics.count_heading_score(paragraph, self.document_parser)
        if heading_score >= 4:
            return 1

        return None
            
    def heading_prefix(self, heading_lvl: int) -> str:
        return "#" * heading_lvl + " "

    def convert_paragraph(self, paragraph: Paragraph) -> str:
        text = ""
        
        sub_parts = list(paragraph.iter_inner_content())

        for sub_part in sub_parts:
            if isinstance(sub_part, str):
                text += sub_part
            elif isinstance(sub_part, Run):
                text += self.convert_paragraph(sub_part)

        return text

def parse_docx(file_name: str) -> str:
    document_parser = DocumentParser(file_name)
    document_statistics = DocumentAnalyzer()
    convertor = MarkdownConverter(document_parser)

    text = convertor.convert_to_markdown(document_statistics)

    return text